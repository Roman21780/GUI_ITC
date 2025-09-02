import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import pyodbc
from db_access import AccessDatabase

import win32com.client
from django.contrib.messages import success
from openpyxl import load_workbook
from openpyxl.cell.cell import MergedCell
import sys
from utils import logger, logging
from openpyxl.styles import numbers

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from datetime import datetime
from docx import Document
import logging
import time
import pythoncom
import unicodedata

import win32timezone  # noqa: F401

import json
import numpy as np
import os
import re
import fitz
import cv2
from docx.shared import Inches
import locale

locale.setlocale(locale.LC_ALL, 'ru_RU.UTF-8')


def format_units(text):
    """Форматирует единицы измерения"""
    return (
        text.replace("г/см3", "г/см³")
            .replace("кгс/см2", "кгс/см²")
            .replace("м2", "м²")
            .replace("м3", "м³")
    )


def superscript(number):
    """Преобразует число в надстрочный формат"""
    superscript_map = {
        "0": "⁰", "1": "¹", "2": "²", "3": "³", "4": "⁴",
        "5": "⁵", "6": "⁶", "7": "⁷", "8": "⁸", "9": "⁹", "-": "⁻"
    }
    return "".join(superscript_map.get(digit, "") for digit in str(number))


def calculate_r_difference(workbook_path, sheet_name='current'):
    """
    Находит разницу между последним значением в столбце R и значением за сутки до последней даты.

    :param workbook_path: путь к файлу Excel
    :param sheet_name: имя листа (по умолчанию 'current')
    :return: разница значений или None в случае ошибки
    """
    import pandas as pd
    from datetime import datetime, timedelta
    import logging

    # Настройка логирования
    logging.basicConfig(level=logging.INFO)

    try:
        # Читаем данные из Excel
        df = pd.read_excel(workbook_path, sheet_name=sheet_name, header=None)

        # Проверяем значение в ячейке B12
        b12_value = df.iloc[11, 1]  # Индексация с 0: строка 12, столбец B
        if isinstance(b12_value, (int, float)) and b12_value < 30:
            logging.info("Значение в ячейке B12 меньше 30. Возвращаем None.")
            return None

        # Получаем столбцы Q и R (индексация с 0, Q=16, R=17)
        column_q = df.iloc[:, 16]  # Столбец Q с датами
        column_r = df.iloc[:, 17]  # Столбец R со значениями

        # 1. Находим последнее значение в столбце R
        last_r_value = column_r.dropna().iloc[-1]

        # 2. Находим последнюю дату в столбце Q
        last_date_str = column_q.dropna().iloc[-1]

        # Преобразуем строку даты в datetime объект
        try:
            last_date = datetime.strptime(str(last_date_str), "%d.%m.%Y %H:%M:%S")
        except ValueError:
            try:
                last_date = datetime.strptime(str(last_date_str), "%d.%m.%Y")
            except ValueError:
                logging.error(f"Невозможно преобразовать дату: {last_date_str}")
                return None

        # Вычитаем ровно одни сутки
        previous_date = last_date - timedelta(days=1)

        # Находим значение в столбце R для даты (previous_date)
        previous_r_value = None
        time_tolerance = timedelta(minutes=60)  # допустимая погрешность во времени

        for i in range(len(column_q)):
            current_date_str = str(column_q.iloc[i])
            if not current_date_str or pd.isna(current_date_str):
                continue

            # Пробуем разные форматы даты
            try:
                current_date = datetime.strptime(current_date_str, "%d.%m.%Y %H:%M:%S")
            except ValueError:
                try:
                    current_date = datetime.strptime(current_date_str, "%d.%m.%Y")
                except ValueError:
                    logging.warning(f"Невозможно преобразовать дату: {current_date_str}")
                    continue

            # Ищем значение, которое отличается ровно на 1 день (± допустимую погрешность)
            if (last_date - current_date >= timedelta(days=1) - time_tolerance and
                    last_date - current_date <= timedelta(days=1) + time_tolerance):
                previous_r_value = column_r.iloc[i]
                break

        if previous_r_value is None:
            logging.warning("Не найдено значение за сутки до последней даты или длительность менее 24ч")
            return None

        # 4. Вычисляем разницу
        difference = last_r_value - previous_r_value
        logging.info(f"Разница значений: {difference} (последнее: {last_r_value}, за сутки до: {previous_r_value})")

        return difference

    except Exception as e:
        logging.warning(f"Ошибка при расчете разницы давления за сутки: {str(e)}")
        return None


def clean_text(text):
    """
    Удаляет суррогатные символы из текста.
    """
    if isinstance(text, str):
        return text.encode('utf-8', errors='ignore').decode('utf-8')
    return text


def table_prev_path(filename):
    """Путь к файлам в папке table_prev"""
    if hasattr(sys, '_MEIPASS'):
        base_path = os.path.normpath(sys._MEIPASS)
    else:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, "table_prev", filename)


def templates_path(filename):
    """Путь к файлам в папке templates"""
    if hasattr(sys, '_MEIPASS'):
        base_path = os.path.normpath(sys._MEIPASS)
    else:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, "templates", filename)

def resource_path(relative_path):
    """
    Возвращает путь к ресурсу, учитывая работу через PyInstaller.
    """
    if hasattr(sys, '_MEIPASS'):
        base_path = os.path.normpath(sys._MEIPASS)
    else:
        # base_path = os.path.dirname(os.path.abspath(__file__))
        base_path = os.path.abspath(".")
    full_path = os.path.join(base_path, relative_path)
    print(f"Resource path: {full_path}")  # Логирование пути
    return full_path


def set_font_size(paragraph_or_run, size):
    """Устанавливает размер шрифта для параграфа или отдельного run"""
    if hasattr(paragraph_or_run, 'runs'):
        # Это Paragraph - применяем ко всем runs
        for run in paragraph_or_run.runs:
            run.font.size = Pt(size)
    else:
        # Это Run - применяем непосредственно
        paragraph_or_run.font.size = Pt(size)


# Загрузка текстовых шаблонов
with open(resource_path('text_templates.json'), 'r', encoding='utf-8') as f:
    TEXT_TEMPLATES = json.load(f)


# Функция вставки diagnostic_text на место метки {{diagnostic_text}}
def insert_diagnostic_text(doc, diagnostic_text):
    """Вставляет диагностический текст на место метки {{diagnostic_text}}"""
    # Получаем текст из шаблонов
    diagnostic_content = get_nested_value(TEXT_TEMPLATES, diagnostic_text.split('.'))

    # Ищем метку в параграфах
    for paragraph in doc.paragraphs:
        if "{{diagnostic_text}}" in paragraph.text:
            for run in paragraph.runs:
                if "{{diagnostic_text}}" in run.text:
                    run.text = run.text.replace("{{diagnostic_text}}", diagnostic_content)
                    break
            break


# Вспомогательная функция для получения вложенных значений из словаря
def get_nested_value(dictionary, keys):
    for key in keys:
        dictionary = dictionary[key]
    return dictionary


# ищем таблицу 'Протокол результатов исследования'
def find_results_table(doc):
    """
    Находит таблицу 'Протокол результатов исследования' по тексту перед ней
    Возвращает таблицу или None если не найдена
    """
    # Ищем параграф с заголовком
    for i, paragraph in enumerate(doc.paragraphs):
        if "Протокол результатов исследования" in paragraph.text:
            # Ищем следующую после параграфа таблицу
            next_element = paragraph._element.getnext()

            while next_element is not None:
                # Проверяем, является ли элемент таблицей
                if next_element.tag.endswith('tbl'):
                    # Находим индекс таблицы в документе
                    for table_idx, table in enumerate(doc.tables):
                        if table._element == next_element:
                            return doc.tables[table_idx]

                next_element = next_element.getnext()

    return None


def replace_and_format_table(doc, data):
    """Удаляет строки с нулевыми или отрицательными значениями во 2-м столбце таблицы результатов"""
    logger = logging.getLogger(__name__)
    table = find_results_table(doc)

    if not table:
        logger.warning("Таблица 'Протокол результатов исследования' не найдена")
        return

    logger.info(f"Начало обработки таблицы. Всего строк: {len(table.rows)}")
    rows_to_delete = []

    for row_idx, row in enumerate(table.rows):
        if len(row.cells) < 2:
            logger.debug(f"Строка {row_idx} пропущена - менее 2 столбцов")
            continue

        second_cell = row.cells[1]
        cell_value = second_cell.text.strip()
        logger.debug(f"Строка {row_idx}, значение: '{cell_value}'")

        try:
            # Пробуем преобразовать в число
            num_value = float(cell_value.replace(',', '.'))
            if num_value == 0 or num_value < -1000:
                logger.info(f"Найден 0/большое отрицательное значение в строке {row_idx}: {num_value}")
                rows_to_delete.append(row_idx)
        except ValueError:
            # Если не число, проверяем текстовые значения
            if cell_value in ("0", "0.0", "0,00", "-", "-2146826252"):
                logger.info(f"Найден 0/большое отрицательное значение (текст) в строке {row_idx}: '{cell_value}'")
                rows_to_delete.append(row_idx)

    logger.info(f"Найдено строк для удаления: {len(rows_to_delete)}")

    # Удаляем строки в обратном порядке
    for row_idx in sorted(rows_to_delete, reverse=True):
        if row_idx < len(table.rows):
            logger.debug(f"Удаление строки {row_idx}")
            table._tbl.remove(table.rows[row_idx]._tr)
        else:
            logger.warning(f"Индекс строки {row_idx} вне диапазона")

    logger.info(f"Удалено строк: {len(rows_to_delete)}. Осталось строк: {len(table.rows)}")



# Функция замены меток (без форматирования единиц измерения)
def replace_tags_only(doc, data):
    """Простая замена меток в тексте и таблицах без изменения структуры"""
    # Обработка обычного текста
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
                set_font_size(paragraph, 12)

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in data.items():
                        original_text = paragraph.text
                        if key in original_text:
                            updated_text = original_text.replace(key, str(value))
                            paragraph.text = updated_text
                            # Проверяем, является ли метка единственным содержимым ячейки
                            if original_text.strip() == key:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            else:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            set_font_size(paragraph, 12)


# Функция восстановления единиц измерения
def fix_units(doc):
    """
    Автоматически находит все '2' и '3' после единиц (кгс, г, м)
    и поднимает их в надстрочный индекс, даже если написано раздельно.
    Примеры:
    - кгс/см2 → кгс/см²
    - кгс см2 → кгс см²
    - м2 → м²
    - г/см3 → г/см³
    """
    # Регулярное выражение для поиска всех "2" и "3" после единиц
    pattern = re.compile(r'(кгс|г|м)([/ ]?[см]?)(2|3)')

    def process_paragraph(paragraph):
        for run in paragraph.runs:
            # Заменяем все вхождения в тексте
            text = pattern.sub(lambda m: f"{m.group(1)}{m.group(2)}{'²' if m.group(3) == '2' else '³'}", run.text)
            if text != run.text:
                run.text = text

    # Обрабатываем весь документ
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)


def replace_plain_tags(doc, data):
    """Заменяет метки БЕЗ скобок с сохранением форматирования"""

    # Обработка обычного текста
    for paragraph in doc.paragraphs:
        _process_paragraph_plain(paragraph, data)

    # Обработка таблиц
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    _process_paragraph_plain(
                        paragraph,
                        data,
                        is_table=True,
                        col_idx=col_idx
                    )


def _process_paragraph_plain(paragraph, data, is_table=False, col_idx=0):
    """Обрабатывает параграф с метками без скобок"""
    # Пропускаем параграфы с графиками
    if any("{{Picture" in run.text for run in paragraph.runs):
        return

    # Объединяем Runs для поиска разбитых меток
    full_text = ''.join(run.text for run in paragraph.runs)

    # Заменяем метки
    changed = False
    for tag, value in data.items():
        if tag in full_text:
            full_text = full_text.replace(tag, str(value))
            changed = True

    if not changed:
        return

    # Сохраняем позиции надстрочных символов
    sup_chars = {'²', '³', '⁴'}
    sup_positions = [i for i, c in enumerate(full_text) if c in sup_chars]

    # Восстанавливаем текст
    paragraph.clear()
    new_run = paragraph.add_run(full_text)

    # Восстанавливаем надстрочные символы
    for pos in sup_positions:
        if pos < len(full_text) and full_text[pos] in sup_chars:
            new_run.font.superscript = True

    # Форматирование для таблиц
    if is_table:
        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.LEFT if col_idx == 0
            else WD_PARAGRAPH_ALIGNMENT.CENTER
        )


def replace_tags_preserve_format(doc, data):
    """Заменяет метки, сохраняя ИСХОДНОЕ форматирование текста"""

    # Обработка обычных параграфов
    for paragraph in doc.paragraphs:
        _process_paragraph_preserve(paragraph, data)

    # Обработка таблиц
    for table in doc.tables:
        for row in table.rows:
            for col_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    _process_paragraph_preserve(
                        paragraph,
                        data,
                        is_table=True,
                        col_idx=col_idx
                    )


def _process_paragraph_preserve(paragraph, data, is_table=False, col_idx=0):
    """Обрабатывает параграф с полным сохранением форматирования"""
    # Пропускаем параграфы с графиками
    if any("{{Picture" in run.text for run in paragraph.runs):
        return

    # Работаем с каждым Run отдельно
    for run in paragraph.runs:
        original_text = run.text
        if not any(tag in original_text for tag in data.keys()):
            continue

        # Сохраняем ВСЕ атрибуты форматирования
        original_font = {
            'name': run.font.name,
            'size': run.font.size,
            'bold': run.font.bold,
            'italic': run.font.italic,
            'underline': run.font.underline,
            'color': run.font.color.rgb if run.font.color else None,
            'superscript': run.font.superscript,
            'subscript': run.font.subscript
        }

        # Заменяем метки
        new_text = original_text
        for tag, value in data.items():
            if tag in original_text:
                new_text = new_text.replace(tag, format_units(str(value)))

        # Применяем изменения
        run.text = new_text

        # Восстанавливаем ВСЕ атрибуты
        run.font.name = original_font['name']
        run.font.size = original_font['size']  # Сохраняем исходный размер шрифта
        run.font.bold = original_font['bold']
        run.font.italic = original_font['italic']
        run.font.underline = original_font['underline']
        if original_font['color']:
            run.font.color.rgb = original_font['color']
        run.font.superscript = original_font['superscript']  # Восстанавливаем надстрочные
        run.font.subscript = original_font['subscript']

    # Выравнивание для таблиц
    if is_table:
        paragraph.alignment = (
            WD_PARAGRAPH_ALIGNMENT.LEFT if col_idx == 0
            else WD_PARAGRAPH_ALIGNMENT.CENTER
        )



def fix_split_runs(paragraph):
    """Объединяет разбитые Runs (если метка разорвана)"""
    if len(paragraph.runs) <= 1:
        return

    full_text = ''.join(run.text for run in paragraph.runs)
    paragraph.clear()
    new_run = paragraph.add_run(full_text)
    # Переносим форматирование первого Run
    if paragraph.runs:
        first_run = paragraph.runs[0]
        first_run.font.size = Pt(12)


def normalize_text(text):
    """
    Нормализует текст, удаляя или заменяя специальные символы.
    """
    if isinstance(text, str):
        # Нормализация текста
        normalized_text = unicodedata.normalize('NFKD', text)
        # Удаление всех символов, которые не являются буквами, цифрами или пробелами
        return ''.join(c for c in normalized_text if unicodedata.category(c) != 'So')
    return text


def safe_quit_office(app, visible_status=None):
    """Безопасное закрытие офисного приложения только если мы его создавали"""
    try:
        if app and hasattr(app, 'Quit') and (visible_status is None or not visible_status):
            app.Quit()
    except Exception as e:
        logger.warning(f"Ошибка при закрытии офисного приложения: {str(e)}")


def normalize_string(value):
    if isinstance(value, str):
        # Если есть символы степени — не трогаем
        if any(c in value for c in ['²', '³', '⁴', '⁵', '⁻']):
            return value

        value = value.strip()
        value = re.sub(r'\s+', ' ', value)
        value = value.replace('\xa0', ' ')
        value = value.replace('\n', ' ').replace('\r', '')
    return value


def extract_numbers_before_letter(value):
    match = re.match(r'(\d+)', value)
    return match.group(0) if match else ''


def convert_to_datetime(value):
    if not value or value == '-':
        return None
    if isinstance(value, str):
        try:
            return datetime.strptime(value, "%d.%m.%Y")
        except ValueError:
            return None
    elif isinstance(value, datetime):
        return value
    return None


def copy_excel_to_word_pandas(excel_path, word_path, sheet_name, search_text):
    """Копирует данные из Excel (все 40 столбцов) в Word документ, заменя указанную метку таблицей."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import pandas as pd
    from docx import Document
    import win32com.client
    import pythoncom
    import os
    import logging

    # Настройка логирования
    logging.basicConfig(level=logging.INFO)

    # 1. Проверка формата через ячейку AM1 и чтение данных через win32com
    pythoncom.CoInitialize()
    use_minimal_columns = False
    all_data = []
    try:
        excel_app = win32com.client.Dispatch("Excel.Application")
        excel_app.Visible = False
        workbook = excel_app.Workbooks.Open(os.path.abspath(excel_path))
        sheet = workbook.Sheets(sheet_name)

        # Проверка формата
        AM1_value = sheet.Range("AM1").Value
        use_minimal_columns = str(AM1_value).strip() == ""
        logging.info(f"Формат таблицы: {'минимальный' if use_minimal_columns else 'полный'}")

        # Чтение всех 40 столбцов (A-AN) и первых 16 строк
        for row in range(1, 17):  # Строки 1-16 (нумерация с 1)
            row_data = []
            for col in range(1, 41):  # Колонки A-AN (1-40)
                cell_value = sheet.Cells(row, col).Value
                row_data.append(cell_value)
            all_data.append(row_data)

    except Exception as e:
        logging.error(f"Ошибка чтения Excel через win32com: {str(e)}")
        return None
    finally:
        # Гарантированное освобождение ресурсов
        if workbook:
            try:
                workbook.Close(False)  # Закрываем только книгу
            except Exception as e:
                logging.warning(f"Ошибка при закрытии workbook в copy_excel_to_word_pandas: {e}")

        excel_app = None
        workbook = None
        # pythoncom.CoUninitialize()

    # 2. Обработка данных
    try:
        # Создаем DataFrame из прочитанных данных
        df = pd.DataFrame(all_data)

        # Проверка структуры данных
        if len(df.columns) < 40:
            logging.warning(f"Файл содержит только {len(df.columns)} из 40 столбцов")

        # Выбор формата таблицы
        if use_minimal_columns:
            columns = [34, 35, 36, 37]  # AI-AL (индексы 34-37, так как в Python нумерация с 0)
        else:
            columns = list(range(34, 40))  # AI-AN (34-39)

        # Фильтрация данных (берем строки 3-16, так как в all_data строки 1-16)
        data_df = df.iloc[2:16, columns].copy()
        data_df.columns = [str(df.iloc[1, col]) for col in columns]

        # Обработка данных
        if not data_df.empty:
            # Первая колонка - дата
            data_df.iloc[:, 0] = pd.to_datetime(data_df.iloc[:, 0], errors='coerce').dt.strftime('%d.%m.%Y')

            # Остальные колонки - числа
            for col in data_df.columns[1:]:
                data_df[col] = pd.to_numeric(data_df[col], errors='coerce').round(1)

        data_df = data_df.dropna(how='all')
        logging.info(f"Данные для вставки:\n{data_df.to_string()}")

    except Exception as e:
        logging.error(f"Ошибка обработки данных: {str(e)}")
        return None

    # 3. Вставка таблицы в Word (остается без изменений)
    try:
        doc = Document(word_path)
        found = False

        for paragraph in doc.paragraphs:
            if search_text in paragraph.text:
                table = doc.add_table(rows=data_df.shape[0] + 1, cols=data_df.shape[1])
                table.style = 'Table Grid'

                # Заголовки
                for col_idx, header in enumerate(data_df.columns):
                    cell = table.cell(0, col_idx)
                    cell.text = str(header)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].font.bold = True

                # Данные
                for row_idx in range(data_df.shape[0]):
                    for col_idx in range(data_df.shape[1]):
                        value = data_df.iat[row_idx, col_idx]
                        cell = table.cell(row_idx + 1, col_idx)
                        cell.text = str(value) if not pd.isna(value) else ''
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Замена метки
                paragraph.text = ''
                paragraph._element.addnext(table._element)
                found = True
                break

        if not found:
            raise ValueError(f"Метка '{search_text}' не найдена в документе")

        output_path = os.path.join(os.path.dirname(word_path), 'KVD_For_Killing.docx')
        doc.save(output_path)
        logging.info(f"Документ сохранен: {output_path}")
        return output_path

    except Exception as e:
        logging.error(f"Ошибка работы с Word: {str(e)}")
        return None


# --------------------------------------------------------------------------------------------------------------



# PDFReader-----------------------------------------------------------------------------------
class PDFReader:
    def __init__(self, pdf_path, output_dir="plots"):
        """
        Инициализация класса для работы с PDF.
        :param pdf_path: Путь к PDF-файлу.
        :param output_dir: Папка для сохранения изображений.
        """
        self.pdf_path = pdf_path
        self.output_dir = output_dir
        self.pages = []
        self.result_list = []

        # Создаем папку для графиков, если её нет
        os.makedirs(self.output_dir, exist_ok=True)

        # Очищаем папку plots перед обработкой
        self.clear_plots_directory()

        # Загружаем страницы PDF
        self.load_pdf()

    def clear_plots_directory(self):
        """
        Удаляет все файлы из папки plots.
        """
        try:
            for file_name in os.listdir(self.output_dir):
                file_path = os.path.join(self.output_dir, file_name)
                if os.path.isfile(file_path):
                    os.remove(file_path)
            logging.info(f"Папка {self.output_dir} успешно очищена.")
        except Exception as e:
            logging.error(f"Ошибка при очистке папки {self.output_dir}: {str(e)}")

    def load_pdf(self):
        """
        Загружает все страницы PDF.
        """
        try:
            doc = fitz.open(self.pdf_path)
            self.pages = [page for page in doc]
        except Exception as e:
            print(f"Ошибка при загрузке PDF: {e}")

    def extract_text(self):
        """
        Извлекает текст со всех страниц PDF.
        """
        for page in self.pages:
            text = page.get_text("text")
            self.result_list.append(text)

    def plot(self, inner_type_page, page_number):
        """
        Сохраняет изображение с указанной страницы PDF.
        Если график встречается дважды, добавляет уникальный суффикс к имени файла.
        """
        trans = {
            "Граф": "Graph",
            "Полулогарифмический": "Semi",
            "µгаза": "Mug",
            "Диагностический": "Log",
            "График": "History",
            "Карта": "Map",
            "Аса": "АСА",
        }

        if inner_type_page in trans:
            inner_type_page = trans[inner_type_page]

        # Генерируем уникальное имя файла с учетом возможных дубликатов
        base_name = f"cropped_image_{inner_type_page}"
        suffix = 1
        while True:
            output_name = f"{base_name}_{suffix}.png"  # Добавляем суффикс
            output_path = os.path.join(self.output_dir, output_name)
            if not os.path.exists(output_path):  # Проверяем, существует ли файл
                break
            suffix += 1

        # Открываем PDF и сохраняем изображение
        doc = fitz.open(self.pdf_path)
        page = doc.load_page(page_number)
        pix = page.get_pixmap(dpi=200)
        pix.save(output_path)

        # Обрезаем изображение
        image = cv2.imdecode(np.fromfile(output_path, dtype=np.uint8), cv2.IMREAD_UNCHANGED)
        cropped_image = image[300:1100, 150:1500]
        cv2.imwrite(output_path, cropped_image)

        return output_path

    def process_pdf(self):
        """
        Основной метод для обработки PDF.
        """
        self.extract_text()
        for page_number, page_text in enumerate(self.result_list):
            type_page_reg = r'^\w+'
            type_page = re.search(type_page_reg, page_text)
            if type_page:
                graphic = self.plot(type_page.group(), page_number)
                print(f"Сохранено изображение: {graphic}")




# GUI--------------------------------------------------------------------------------

def ensure_python_dll():
    if getattr(sys, 'frozen', False):
        # Для собранного приложения (PyInstaller)
        base_path = sys._MEIPASS
        dll_path = os.path.join(base_path, 'python312.dll')
        if not os.path.exists(dll_path):
            raise FileNotFoundError(f"Файл python312.dll не найден в {dll_path}!")
    else:
        # Для запуска из исходников
        dll_path = os.path.join(os.path.dirname(__file__), 'python312.dll')
        if not os.path.exists(dll_path):
            raise FileNotFoundError(f"Файл python312.dll не найден в {dll_path}!")


# Вызов функции в начале работы программы
# ensure_python_dll()

logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    filename='app.log',  # Имя файла для логов
    filemode='w'  # 'w' для перезаписи файла, 'a' для добавления
)


class ReportGUI:
    def __init__(self, root):
        self.root = root

        # Проверка наличия python312.dll
        if getattr(sys, 'frozen', False):
            sys._MEIPASS = os.path.normpath(sys._MEIPASS)

        self.pdf_path = None
        self.output_file_path = None
        self.root.title("Параметры проекта")
        self.root.minsize(900, 600)

        self.color_buttons = []

        self.insert_button = None
        self.model_button = None
        self.pressure_button = None
        self.model_button2 = None
        self.pressure_button2 = None
        self.insert_button2 = None
        self.insert_button2_2 = None

        # Создаем основной фрейм с прокруткой
        self.main_frame = ttk.Frame(self.root)
        self.main_frame.pack(fill='both', expand=True)

        # Создаем canvas и scrollbar
        self.canvas = tk.Canvas(self.main_frame)
        self.scrollbar = ttk.Scrollbar(self.main_frame, orient="vertical", command=self.canvas.yview)
        self.scrollable_frame = ttk.Frame(self.canvas)

        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )

        self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)

        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        # Привязываем прокрутку колесом мыши
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        self.scrollable_frame.bind("<Enter>", self._bind_mousewheel)
        self.scrollable_frame.bind("<Leave>", self._unbind_mousewheel)

        # Инициализация стиля для зеленой кнопки
        self.style = ttk.Style()
        self.style.configure("Green.TButton", background="green", foreground="green")

        self.base_dir = os.path.dirname(__file__)

        # Инициализация базы данных
        self.db = AccessDatabase()
        self.current_main_data_id = None

        self.section_params = {
            1: {"start_cell": "A1", "expected_columns": 2, "description": "Входные данные"},
            2: {"start_cell": "L1", "expected_columns": 4, "description": "Модель давления на ВНК"},
            3: {"start_cell": "Q1", "expected_columns": 2, "description": "Данные давления на ВНК"},
            4: {"start_cell": "G1", "expected_columns": 3, "description": "Параметры_2"},
            5: {"start_cell": "T1", "expected_columns": 4, "description": "Модель давления_2"},
            6: {"start_cell": "Y1", "expected_columns": 2, "description": "Данные давления_2"}
        }

        # Очистка базы данных при запуске
        self.clear_database_on_startup()

        self.setup_gui(self.scrollable_frame)

    def _on_canvas_configure(self, event):
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def _bind_mousewheel(self, event):
        self.canvas.bind_all("<MouseWheel>", self._on_mousewheel)

    def _unbind_mousewheel(self, event):
        self.canvas.unbind_all("<MouseWheel>")

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def clear_database_on_startup(self):
        """Очищает базу данных при запуске приложения"""
        try:
            success = self.db.clear_data()  # Теперь этот метод существует
            if success:
                logging.info("База данных очищена при запуске")
            else:
                logging.error("Не удалось очистить базу данных")
        except Exception as e:
            logging.error(f"Ошибка при очистке базы данных: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при очистке базы данных: {str(e)}")

    def clear_database(self):
        """Очищает базу данных и сбрасывает цвет кнопок"""
        try:
            success = self.db.clear_data()
            if success:
                self.reset_button_colors()
                messagebox.showinfo("Успех", "Данные успешно удалены из базы данных")
                return True
            else:
                messagebox.showerror("Ошибка", "Не удалось очистить базу данных")
                return False
        except Exception as e:
            logging.error(f"Ошибка при очистке базы данных: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при очистке базы данных: {str(e)}")
            return False

    def reset_button_colors(self):
        """Сбрасывает цвет всех кнопок, которые были зелеными"""
        for button in self.color_buttons:
            button.config(style="TButton")
        self.color_buttons = []

    def on_close(self):
        """Обработчик закрытия приложения"""
        if hasattr(self, 'db'):
            self.db.close()
        self.root.destroy()

    def setup_gui(self, parent):
        self.notebook = ttk.Notebook(parent)
        self.notebook.pack(fill='both', expand=True, padx=5, pady=5)

        self.tab1 = ttk.Frame(self.notebook)
        self.tab2 = ttk.Frame(self.notebook)
        self.notebook.add(self.tab1, text='Основные данные')
        self.notebook.add(self.tab2, text='Давление')

        self.setup_tab1()
        self.setup_tab2()
        self.setup_bottom_buttons()
        self.setup_pdf_processing()  # Добавляем обработку PDF
        # Добавляем обработчик закрытия окна
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)

    def create_labeled_frame(self, parent, text):
        frame = ttk.LabelFrame(parent, text=text)
        frame.pack(fill='x', padx=3, pady=3)
        return frame

    def setup_tab1(self):
        # Фрейм для ввода данных
        frame1 = self.create_labeled_frame(self.tab1, f"1. {self.section_params[1]['description']}")
        ttk.Label(frame1, text="Скопируйте данные:").pack(padx=3, pady=3)
        self.insert_button = ttk.Button(frame1, text="Вставить", command=lambda: self.paste_data(1))
        self.insert_button.pack(padx=3, pady=3)

        # Класс исследования
        frame2 = self.create_labeled_frame(self.tab1, "2. Класс исследования")
        ttk.Label(frame2, text="Класс (1-3):\n 1 - Рпл\n 2 - Кпрод (Рпл+Рзаб)\n 3 - Рпл, ФЕС, Кпрод").pack(padx=3,
                                                                                                           pady=3)
        self.class_entry = ttk.Entry(frame2)
        self.class_entry.pack(padx=3, pady=3)

        # Успешность исследования
        frame3 = self.create_labeled_frame(self.tab1, "3. Успешность исследования")
        ttk.Label(frame3,
                  text="Успешность (1-7):\n 1 - тех.неиспр(отказ, дрейф, шум)\n 2 - уход уровня\n 3 - АРД\n 4 - ЗКЦ\n 5 - НД\n 6 - пропуски ФА\n 7 - границы/интерф").pack(
            padx=3, pady=3)
        self.success_entry = ttk.Entry(frame3)
        self.success_entry.pack(padx=3, pady=3)

        # Поправки Pпл, Pзаб, Pпл_2, Pзаб_2
        correction_frame = self.create_labeled_frame(self.tab1, "4. Поправки")
        self.setup_correction_fields(correction_frame)

        # Плотность
        frame8 = self.create_labeled_frame(self.tab1, "5. Плотность")
        ttk.Label(frame8, text="пересчет Pзаб, (г/см3):").pack(padx=3, pady=3)
        self.density_zab_entry = ttk.Entry(frame8)
        self.density_zab_entry.pack(padx=3, pady=3)
        ttk.Label(frame8, text="пересчет Pпл, (г/см3):").pack(padx=3, pady=3)
        self.density_pl_entry = ttk.Entry(frame8)
        self.density_pl_entry.pack(padx=3, pady=3)

        # Вставка параметров исследования
        params_frame = self.create_labeled_frame(self.tab1, "6. Параметры исследования")

        self.insert_button2 = ttk.Button(
            params_frame,
            text="Вставить параметры",
            command=self.paste_research_params  # ← Без скобок!
        )
        self.insert_button2.pack(padx=3, pady=3)

        self.insert_button2_2 = ttk.Button(
            params_frame,
            text="Вставить параметры_2",
            command=self.paste_research_params_2  # ← Без скобок!
        )
        self.insert_button2_2.pack(padx=3, pady=3)

        # Добавляем поле для "Расчетное время"
        frame_time = ttk.Frame(self.tab1)
        frame_time.pack(fill='x', padx=5, pady=5)

        ttk.Label(frame_time, text="Расчетное время, ч:").pack(side='left', padx=5)
        self.calc_time_entry = ttk.Entry(frame_time)
        self.calc_time_entry.pack(side='left', padx=5, fill='x', expand=False)

        # Добавляем поле для давления на последнюю точку
        pressure_frame = self.create_labeled_frame(self.tab1, "7. Давление на последнюю точку")
        ttk.Label(pressure_frame, text="Р_п.т.:").pack(padx=3, pady=3)
        self.pressure_last_point_entry = ttk.Entry(pressure_frame)
        self.pressure_last_point_entry.pack(padx=3, pady=3)


    def change_button_color(self, button, success=True):
        """Изменяет цвет кнопки на зеленый (успех) или сбрасывает (ошибка)."""
        if success:
            button.config(style="Green.TButton")
            # Добавляем кнопку в список, если её там еще нет
            if button not in self.color_buttons:
                self.color_buttons.append(button)
        else:
            # Сбрасываем стиль кнопки
            button.config(style="TButton")

    def setup_correction_fields(self, parent):
        # Создаем основной фрейм для размещения всех полей
        main_frame = ttk.Frame(parent)
        main_frame.pack(fill='both', expand=True, padx=5, pady=5)

        # Основной фрейм для основных поправок (4 колонки)
        corrections_frame = ttk.Frame(main_frame)
        corrections_frame.pack(side='top', fill='x', pady=5)

        # Определяем заголовки для поправок
        correction_labels = [
            "Поправка на ВНК",
            "Поправка на ВДП",
            "Поправка на ГНК"
        ]

        # Создаем четыре колонки для поправок
        columns = []
        for i in range(4):
            column_frame = ttk.Frame(corrections_frame)
            column_frame.pack(side='left', fill='y', padx=10, pady=5)
            columns.append(column_frame)

        # Заголовки для каждой колонки
        column_titles = ["Pпл", "Pзаб", "Pпл_2", "Pзаб_2"]

        # Создаем поля для каждой колонки
        self.ppl_entries = []
        self.pzab_entries = []
        self.ppl2_entries = []
        self.pzab2_entries = []

        entry_lists = [
            self.ppl_entries,
            self.pzab_entries,
            self.ppl2_entries,
            self.pzab2_entries,
        ]

        for i, column_frame in enumerate(columns):
            # Добавляем заголовок для колонки
            ttk.Label(column_frame, text=column_titles[i], font=("Arial", 10, "bold")).pack(anchor='w', padx=3, pady=3)

            # Добавляем поля для поправок
            for label in correction_labels:
                frame = ttk.Frame(column_frame)
                frame.pack(fill='x', padx=3, pady=2)
                ttk.Label(frame, text=f"{label}:").pack(side='left', padx=3)
                entry = ttk.Entry(frame, width=10)
                entry.pack(side='left', padx=3)
                entry_lists[i].append(entry)

        # Поправки на ВНК пластов 3 и 4
        # Добавляем отдельные поля для "Поправка на ВНК Рпл_3" и "Поправка на ВНК Рпл_4"
        additional_frame = ttk.Frame(main_frame)
        additional_frame.pack(side='top', fill='x', padx=5, pady=5)

        ttk.Label(additional_frame, text="Поправка на ВНК Рпл_3:", font=("Arial", 8)).pack(side='left', padx=5)
        self.vnkp_pl3_entry = ttk.Entry(additional_frame, width=8, font=("Arial", 8))
        self.vnkp_pl3_entry.pack(side='left', padx=5)

        ttk.Label(additional_frame, text="Поправка на ВНК Рпл_4:", font=("Arial", 8)).pack(side='left', padx=5)
        self.vnkp_pl4_entry = ttk.Entry(additional_frame, width=8, font=("Arial", 8))
        self.vnkp_pl4_entry.pack(side='left', padx=5)

    def setup_tab2(self):
        frame6 = self.create_labeled_frame(self.tab2, f"8. {self.section_params[2]['description']}")

        # NEW: Кнопки с отслеживанием нажатия
        self.model_button = ttk.Button(
            frame6,
            text="Модель",
            command=lambda: [self.paste_data(2), self.change_button_color(self.model_button)]
        )
        self.model_button.pack(padx=3, pady=3)

        self.model_button2 = ttk.Button(
            frame6,
            text="Модель_2",
            command=lambda: [self.paste_data(5), self.change_button_color(self.model_button2)]
        )
        self.model_button2.pack(padx=3, pady=3)

        frame7 = self.create_labeled_frame(self.tab2, f"9. {self.section_params[3]['description']}")

        self.pressure_button = ttk.Button(
            frame7,
            text="Давление",
            command=lambda: [self.paste_data(3), self.change_button_color(self.pressure_button)]
        )
        self.pressure_button.pack(padx=3, pady=3)

        self.pressure_button2 = ttk.Button(
            frame7,
            text="Давление_2",
            command=lambda: [self.paste_data(6), self.change_button_color(self.pressure_button2)]
        )
        self.pressure_button2.pack(padx=3, pady=3)

    def paste_research_params(self):
        """Вставляет параметры исследования из буфера обмена"""
        try:
            clipboard_data = self.root.clipboard_get()
            if not clipboard_data.strip():
                messagebox.showerror("Ошибка", "Буфер обмена пуст.")
                return

            main_data_id = self.ensure_main_data_exists()
            if main_data_id is None:
                messagebox.showerror("Ошибка", "Не удалось создать основную запись!")
                return

            # Используем новый метод для трехколоночных данных
            self.db.insert_all_calculated_parameters_from_clipboard(main_data_id, clipboard_data)
            messagebox.showinfo("Успех", "Параметры исследования сохранены успешно")

        except Exception as e:
            logging.error(f"Ошибка при вставке параметров исследования: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при вставке параметров: {str(e)}")

    def paste_research_params_2(self):
        """Вставляет параметры исследования_2 из буфера обмена"""
        try:
            clipboard_data = self.root.clipboard_get()
            if not clipboard_data.strip():
                messagebox.showerror("Ошибка", "Буфер обмена пуст.")
                return

            rows = [r.split('\t') for r in clipboard_data.split('\n') if r.strip()]
            if not rows:
                messagebox.showerror("Ошибка", "Нет данных для вставки")
                return

            params = {}
            for row in rows:
                if len(row) >= 2:
                    param_name = row[0].strip()
                    param_value = row[1].strip()
                    if param_name and param_value:
                        converted_value = ReportGUI.convert_value(param_value)
                        if converted_value is not None:
                            params[param_name] = converted_value

            main_data_id = self.ensure_main_data_exists()
            if main_data_id is None:
                messagebox.showerror("Ошибка", "Не удалось создать основную запись!")
                return

            self.db.insert_all_calculated_parameters_from_clipboard(main_data_id,params)
            messagebox.showinfo("Успех", "Параметры исследования_2 сохранены успешно")

        except Exception as e:
            logging.error(f"Ошибка при вставке параметров исследования_2: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при вставке параметров: {str(e)}")

    def select_output_directory(self):
        """Выбираем только директорию, без имени файла"""
        directory = filedialog.askdirectory(title="Выберите папку для сохранения отчетов")
        if directory:
            self.output_directory = directory
            messagebox.showinfo("Успех", f"Отчеты будут сохраняться в: {directory}")
            return True
        return False

    def setup_bottom_buttons(self):
        button_frame = ttk.Frame(self.root)
        button_frame.pack(pady=10)

        template_frame = ttk.Frame(button_frame)
        template_frame.pack(fill='x', pady=3)

        ttk.Label(template_frame, text="Шаблон:").pack(side='left', padx=3)
        self.template_var = tk.StringVar()
        self.template_combobox = ttk.Combobox(template_frame, textvariable=self.template_var, state="readonly", width=15)
        self.template_combobox['values'] = (
            "КВД_Заполярка", "КВД_Оренбург", "КВД_Оренбург_газ", "КВД_Оренбург2",
            "КВД_Хантос", "КВД_глушение", "КВД_ННГ", "КВД+ИД", "КСД",
            "КПД", "КПД+ИД", "ГРП"
        )
        self.template_combobox.current(0)
        self.template_combobox.pack(side='left', padx=3)

        clear_btn = ttk.Button(
            button_frame,
            text="Очистить данные",
            command=self.clear_database
        )
        clear_btn.pack(side='right', padx=100)

        self.select_button = ttk.Button(
            button_frame,
            text="Выбрать папку для сохранения",
            command=lambda: [self.select_output_directory(), self.change_button_color(self.select_button)]
        )
        self.select_button.pack(side='left', padx=5)

        self.save_button = ttk.Button(button_frame, text="Сохранить внесенные данные", command=self.save_to_db)
        self.save_button.pack(side='left', padx=5)

        calculate_btn = ttk.Button(button_frame, text="Произвести расчет", command=self.perform_calculations)
        calculate_btn.pack(side='left', padx=5)

        generate_btn = ttk.Button(button_frame, text="Формировать отчет", command=self.generate_report)
        generate_btn.pack(side='left', padx=5)

    def setup_pdf_processing(self):
        """Добавляет элементы интерфейса для обработки PDF."""
        pdf_frame = ttk.LabelFrame(self.tab1, text="Обработка PDF")
        pdf_frame.pack(fill='x', padx=3, pady=3)

        # Поле для выбора PDF-файла
        self.pdf_var = tk.StringVar()
        ttk.Label(pdf_frame, text="PDF файл:").pack(side='left', padx=3)
        ttk.Entry(pdf_frame, textvariable=self.pdf_var, state="readonly", width=30).pack(side='left', padx=3)

        # Кнопка для выбора PDF
        ttk.Button(pdf_frame, text="Выбрать PDF", command=self.select_pdf).pack(side='left', padx=3)

        # Кнопка для обработки PDF
        ttk.Button(pdf_frame, text="Обработать PDF", command=self.process_pdf).pack(side='left', padx=3)

    def select_pdf(self):
        """Выбирает PDF-файл."""
        pdf_path = filedialog.askopenfilename(
            title="Выберите PDF файл",
            filetypes=[("PDF files", "*.pdf")]
        )
        if pdf_path:
            self.pdf_var.set(pdf_path)
            messagebox.showinfo("Успех", f"Выбран PDF файл: {pdf_path}")

    def process_pdf(self):
        """Обрабатывает выбранный PDF файл."""
        pdf_path = self.pdf_var.get()
        if not pdf_path:
            messagebox.showerror("Ошибка", "PDF файл не выбран!")
            return

        try:
            # Создаем экземпляр PDFReader
            pdf_reader = PDFReader(pdf_path)

            # Обрабатываем PDF
            pdf_reader.process_pdf()
            messagebox.showinfo("Успех", "PDF успешно обработан!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при обработке PDF: {str(e)}")

    def log_invalid_characters(self, text):
        """
        Логирует недопустимые символы в тексте.
        """
        invalid_chars = [char for char in text if ord(char) > 65535]
        if invalid_chars:
            print(f"Обнаружены недопустимые символы: {invalid_chars}")

    @normalize_text
    @clean_text
    def paste_data(self, section):
        """Общая функция для вставки данных из буфера обмена"""
        try:
            clipboard_data = self.root.clipboard_get()
            if not clipboard_data.strip():
                messagebox.showerror("Ошибка", "Буфер обмена пуст")
                return False

            # Обрабатываем данные
            rows = [r.split('\t') for r in clipboard_data.split('\n') if r.strip()]

            if section == 1:  # Основные данные
                return self.process_main_data(rows)
            elif section in [2, 5]:  # Модели
                return self.process_model_data(rows, section)
            elif section in [3, 6]:  # Давление
                return self.process_pressure_data(rows, section)
            else:
                return False

        except Exception as e:
            logging.error(f"Ошибка при вставке данных: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при вставке данных: {str(e)}")
            return False

    def clean_text(self, text):
        """Очищает текст от недопустимых символов"""
        if isinstance(text, str):
            return text.encode('utf-8', errors='ignore').decode('utf-8')
        return text

    def get_button_by_section(self, section):
        """Возвращает кнопку по номеру секции"""
        button_map = {
            1: self.insert_button,
            2: self.model_button,
            3: self.pressure_button,
            5: self.model_button2,
            6: self.pressure_button2
        }
        return button_map.get(section)


    def process_parameters_data(self, rows, main_data_id):
        """Обрабатывает параметры исследования"""
        params = {}
        for row in rows:
            if len(row) >= 2:
                param_name = row[0].strip()
                param_value = self.convert_value(row[1].strip())
                if param_name and param_value is not None:
                    params[param_name] = param_value

        try:
            self.db.insert_all_calculated_parameters_from_clipboard(main_data_id, params)
            messagebox.showinfo("Успех", "Параметры исследования сохранены в БД")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении параметров: {str(e)}")

    def paste_calculated_parameters(self):
        """Вставляет calculatedParameters из буфера обмена"""
        try:
            clipboard_data = self.root.clipboard_get()
            if not clipboard_data.strip():
                messagebox.showerror("Ошибка", "Буфер обмена пуст!")
                return

            rows = [r.split('\t') for r in clipboard_data.split('\n') if r.strip()]

            main_data_id = self.ensure_main_data_exists()
            if main_data_id is None:
                messagebox.showerror("Ошибка", "Нет основной записи!")
                return

            params = {}
            for row in rows:
                if len(row) >= 2:
                    param_name = row[0].strip()
                    param_value = row[1].strip()
                    if param_name and param_value:
                        params[param_name] = ReportGUI.convert_value(param_value)

            if params:
                self.db.insert_all_calculated_parameters_from_clipboard(main_data_id, params)
                messagebox.showinfo("Успех", "Параметры сохранены!")
            else:
                messagebox.showinfo("Информация", "Нет данных для сохранения")

        except Exception as e:
            logging.error(f"Ошибка вставки параметров: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")

    def paste_model_ksd(self):
        """Вставляет данные ModelKSD из буфера обмена"""
        try:
            clipboard_data = self.root.clipboard_get()
            if not clipboard_data.strip():
                messagebox.showerror("Ошибка", "Буфер обмена пуст!")
                return

            print(f"Данные из буфера: {clipboard_data[:200]}...")  # Первые 200 символов

            rows = [r.split('\t') for r in clipboard_data.split('\n') if r.strip()]
            print(f"Разобрано строк: {len(rows)}")

            main_data_id = self.ensure_main_data_exists()
            if main_data_id is None:
                messagebox.showerror("Ошибка", "Нет основной записи!")
                return

            data = []
            for i, row in enumerate(rows):
                if i < 5:  # Покажем первые 5 строк
                    print(f"Строка {i}: {row}")

                if len(row) >= 3:  # 3 столбца: empty, Dat, PressureVnkModel
                    # Пропускаем заголовки
                    if any(header in row[0] for header in ['empty', 'Empty', 'Пустой']):
                        continue
                    if any(header in row[1] for header in ['Dat', 'Date', 'Дата']):
                        continue
                    if any(header in row[2] for header in ['Pressure', 'Model', 'Давление']):
                        continue

                    data.append({
                        'empty': row[0].strip(),
                        'Dat': ReportGUI.convert_value(row[1].strip()),
                        'PressureVnkModel': ReportGUI.convert_value(row[2].strip())
                    })

            print(f"Обработано записей: {len(data)}")
            if data:
                if len(data) > 5:
                    print(f"Первые 5 записей: {data[:5]}")
                else:
                    print(f"Все записи: {data}")

                self.db.insert_model_ksd(main_data_id, data)
                messagebox.showinfo("Успех", "Данные ModelKSD сохранены!")
            else:
                messagebox.showinfo("Информация", "Нет данных для сохранения")

        except Exception as e:
            logging.error(f"Ошибка вставки ModelKSD: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")


    def save_to_db(self):
        """Сохраняет данные из полей ввода в соответствующие таблицы"""
        try:
            # Убеждаемся, что есть основная запись
            main_data_id = self.ensure_main_data_exists()
            if main_data_id is None:
                messagebox.showerror("Ошибка", "Не удалось создать основную запись!")
                return

            # ОЧИСТКА: Удаляем старые данные перед сохранением новых
            self.db.clear_related_data(main_data_id)

            # Подготавливаем данные для сохранения
            update_data = {}

            # Класс исследования
            class_value = self.class_entry.get()
            if class_value:
                update_data['Klass_issledovaniya'] = ReportGUI.convert_value(class_value)

            # Успешность исследования
            success_value = self.success_entry.get()
            if success_value:
                update_data['Uspeshnost'] = ReportGUI.convert_value(success_value)

            # Расчетное время
            calc_time = self.calc_time_entry.get()
            if calc_time:
                update_data['Durat'] = ReportGUI.convert_value(calc_time)

            # Плотность
            density_zab = self.density_zab_entry.get()
            if density_zab:
                update_data['density_zab'] = ReportGUI.convert_value(density_zab)

            density_pl = self.density_pl_entry.get()
            if density_pl:
                update_data['density_pl'] = ReportGUI.convert_value(density_pl)

            # Давление на последнюю точку
            pressure_pt = self.pressure_last_point_entry.get()
            if pressure_pt:
                self.db.insert_pressure_last_point(main_data_id, ReportGUI.convert_value(pressure_pt))

            # Поправки из вашего GUI
            amendments_data = self.collect_amendments_from_gui()
            if amendments_data:
                self.db.insert_amendments(main_data_id, amendments_data)

            # Сохраняем данные
            if update_data:
                self.db.update_main_data(main_data_id, update_data)

            messagebox.showinfo("Успех", "Данные успешно сохранены в базу данных")

        except Exception as e:
            logging.error(f"Ошибка при сохранении данных: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при сохранении данных: {str(e)}")


    def collect_amendments_from_gui(self):
        """Собирает поправки из полей ввода GUI"""
        amendments_data = {}

        # Маппинг ваших полей на имена в базе данных
        field_mapping = {
            # Pпл (первая модель)
            (0, 0): 'amendVnkPpl',  # ВНК Рпл
            (0, 1): 'amendVdpPpl',  # ВДП Рпл
            (0, 2): 'amendGnkPpl',  # ГНК Рпл

            # Pзаб (первая модель)
            (1, 0): 'amendVnkPzab',  # ВНК Рзаб
            (1, 1): 'amendVdpPzab',  # ВДП Рзаб
            (1, 2): 'amendGnkPzab',  # ГНК Рзаб

            # Pпл_2 (вторая модель)
            (2, 0): 'amendVnkPpl2',  # ВНК Рпл_2
            (2, 1): 'amendVdpPpl2',  # ВДП Рпл_2
            (2, 2): 'amendGnkPpl2',  # ГНК Рпл_2

            # Pзаб_2 (вторая модель)
            (3, 0): 'amendVnkPzab2',  # ВНК Рзаб_2
            (3, 1): 'amendVdpPzab2',  # ВДП Рзаб_2
            (3, 2): 'amendGnkPzab2',  # ГНК Рзаб_2
        }

        # Собираем данные из всех полей
        all_entries = [
            self.ppl_entries,  # Pпл
            self.pzab_entries,  # Pзаб
            self.ppl2_entries,  # Pпл_2
            self.pzab2_entries  # Pзаб_2
        ]

        for col_idx, entries in enumerate(all_entries):
            for row_idx, entry in enumerate(entries):
                value = entry.get().strip()
                if value:
                    field_name = field_mapping.get((col_idx, row_idx))
                    if field_name:
                        amendments_data[field_name] = ReportGUI.convert_value(value)

        # Дополнительные поправки
        vnk_pl3 = self.vnkp_pl3_entry.get().strip()
        if vnk_pl3:
            amendments_data['amendVnkPpl3'] = ReportGUI.convert_value(vnk_pl3)

        vnk_pl4 = self.vnkp_pl4_entry.get().strip()
        if vnk_pl4:
            amendments_data['amendVnkPpl4'] = ReportGUI.convert_value(vnk_pl4)

        return amendments_data

    def process_main_data(self, rows):
        """Обрабатывает основные данные (секция 1) и сохраняет в базу"""
        try:
            # Жесткий порядок полей (соответствует порядку в буфере обмена)
            field_order = [
                'Company',  # Компания
                'Localoredenie',  # Месторождение
                'Skvazhina',  # Скважина
                'VNK',  # Имя исследования / #
                'Data_issledovaniya',  # Дата исследования
                'Plast',  # Пласт
                'Interval_perforacii',  # Интервал перфорации
                'Tip_pribora',  # Тип прибора / #
                'Glubina_ustanovki_pribora',  # Глубина установки прибора
                'Interpretator',  # Имя интерпретатора
                'Data_interpretacii',  # Дата интерпретации
                'Vremya_issledovaniya',  # Время исследования (без заголовка)
                'Obvodnennost',  # Обводненность (без заголовка)
                'Nalicie_paktera',  # Наличие пакера (без заголовка)
                'Data_GRP',  # Дата ГРП (без заголовка)
                'Vid_issledovaniya'  # Вид исследования (без заголовка)
            ]

            # Собираем все значения из строк
            all_values = []
            for row in rows:
                if len(row) >= 2:
                    value = row[1].strip()
                    if value:  # Добавляем только непустые значения
                        all_values.append(value)

            # Создаем словарь данных
            data_dict = {}
            for i, field_name in enumerate(field_order):
                if i < len(all_values):
                    data_dict[field_name] = ReportGUI.convert_value(all_values[i])

            # Отладочная информация
            print("Данные для сохранения:")
            for key, value in data_dict.items():
                print(f"  {key}: {value}")

            # Сохраняем в базу
            main_data_id = self.db.insert_main_data(data_dict)
            logging.info(f"Основные данные сохранены с ID: {main_data_id}")

            # Сохраняем ID для использования в других методах
            self.current_main_data_id = main_data_id

            messagebox.showinfo("Успех", "Основные данные сохранены в базу данных")
            return main_data_id

        except Exception as e:
            logging.error(f"Ошибка при сохранении основных данных: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при сохранении данных: {str(e)}")
            return None

    def ensure_main_data_exists(self):
        """Проверяет и создает основную запись если нужно"""
        if self.current_main_data_id is None:
            last_record = self.db.get_last_record()
            if not last_record.empty:
                self.current_main_data_id = last_record.iloc[0]['id']
                return self.current_main_data_id
            else:
                # Создаем пустую основную запись
                try:
                    empty_data = {
                        "company": "Не указано",
                        "field": "Не указано",
                        "well": "Не указано",
                        "created_date": datetime.now()
                    }
                    self.current_main_data_id = self.db.insert_main_data(empty_data)
                    logging.info(f"Создана пустая основная запись с ID: {self.current_main_data_id}")
                    return self.current_main_data_id
                except Exception as e:
                    logging.error(f"Ошибка создания пустой записи: {str(e)}")
                    return None
        return self.current_main_data_id

    def process_model_data(self, rows, section):
        """Обрабатывает данные модели и выбирает целевую таблицу"""
        data_list = []
        for row in rows:
            if len(row) >= 3:  # Нужно минимум 3 столбца: empty, Dat, PressureVnkModel
                # Пропускаем заголовки
                if (any(header in row[0] for header in ['empty', 'Empty', 'Пустой']) or
                        any(header in row[1] for header in ['DT', 'Dat', 'Date', 'Дата', 'Прошедшее время', '(ч)']) or
                        any(header in row[2] for header in
                            ['DP', 'deltaP', 'Pressure', 'Model', 'Давление', 'PressureVnkModel', '(кг/см^2 )'])):
                    continue

                # Пропускаем пустые строки
                if not row[0].strip() and not row[1].strip() and not row[2].strip():
                    continue

                data_list.append({
                    'empty': row[0].strip(),
                    'Dat': row[1].strip(),
                    'PressureVnkModel': row[2].strip()
                })

        main_data_id = self.ensure_main_data_exists()
        if main_data_id is None:
            messagebox.showerror("Ошибка", "Сначала вставьте основные данные!")
            return

        try:
            # Получаем тип исследования
            research_type = self.db.get_research_type(main_data_id)
            print(f"Тип исследования: {research_type}")

            if not data_list:
                messagebox.showinfo("Информация", "Нет данных для сохранения после фильтрации заголовков")
                return

            # Определяем целевую таблицу на основе типа исследования
            if research_type and "КСД" in research_type.upper():
                # Для КСД исследований - вставляем в ModelKSD
                print(f"Вставляем {len(data_list)} записей в ModelKSD")
                self.db.insert_model_ksd(main_data_id, data_list)
                messagebox.showinfo("Успех", "Данные модели сохранены в ModelKSD!")
            else:
                # Для всех остальных исследований - вставляем в ModelVNK
                print(f"Вставляем {len(data_list)} записей в ModelVNK")
                self.db.insert_model_vnk(main_data_id, data_list)
                messagebox.showinfo("Успех", "Данные модели сохранены в ModelVNK!")

        except Exception as e:
            logging.error(f"Ошибка при сохранении параметров модели: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка: {str(e)}")


    def process_pressure_data(self, rows, section):
        """Обрабатывает данные давления (секция 3) и сохраняет в PressureVNK"""
        data_list = []
        for row in rows:
            if len(row) >= 2:
                # Пропускаем заголовки
                if row[0].strip() == "Абсолютное время" or row[1].strip() == "кг/см^2":
                    continue

                data_list.append({
                    'Dat': ReportGUI.convert_value(row[0].strip()),
                    'PressureVnk': ReportGUI.convert_value(row[1].strip())
                })

        main_data_id = getattr(self, 'current_main_data_id', None)
        if main_data_id is None:
            last_record = self.db.get_last_record()
            if not last_record.empty:
                main_data_id = last_record.iloc[0]['id']
            else:
                logging.error("Нет основной записи для привязки данных давления")
                messagebox.showerror("Ошибка", "Сначала вставьте основные данные!")
                return

        try:
            # ИСПРАВЛЕНО: вставляем в PressureVNK
            self.db.insert_pressure_vnk_text_file(main_data_id, data_list)
            messagebox.showinfo("Успех", "Данные давления сохранены в базу данных")
        except Exception as e:
            logging.error(f"Ошибка при сохранении данных давления: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при сохранении данных: {str(e)}")


    @staticmethod
    def convert_value(value):
        """Конвертирует строковое значение в соответствующий тип"""
        if not value or value == '' or value == 'None' or value is None:
            return None

        if not isinstance(value, str):
            return value

        value = value.strip()

        # Обработка специальных значений
        if value in ['-', '—', '–', 'н/д', 'не указано']:
            return None

        # Пробуем преобразовать в число
        try:
            value_clean = value.replace(',', '.')
            return float(value_clean)
        except ValueError:
            pass

        # Пробуем преобразовать в целое число
        try:
            return int(value)
        except ValueError:
            pass

        # Пробуем преобразовать в дату (разные форматы)
        try:
            from datetime import datetime
            # Пробуем разные форматы дат
            for fmt in ['%d.%m.%Y', '%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', '%d.%m.%y']:
                try:
                    return datetime.strptime(value, fmt).date()
                except ValueError:
                    continue
        except:
            pass

        # Оставляем как строку
        return value

    def clear_amendments_fields(self):
        """Очищает поля ввода поправок"""
        for entry in self.ppl_entries + self.pzab_entries + self.ppl2_entries + self.pzab2_entries:
            entry.delete(0, 'end')
        self.vnkp_pl3_entry.delete(0, 'end')
        self.vnkp_pl4_entry.delete(0, 'end')

    def select_output_file(self):
        self.output_file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Выберите место сохранения"
        )
        if self.output_file_path:
            messagebox.showinfo("Успех", f"Выбрано место сохранения: {self.output_file_path}")
        else:
            messagebox.showwarning("Предупреждение", "Место сохранения не выбрано.")

    @staticmethod
    def show_meipass_content():
        try:
            if getattr(sys, 'frozen', False):
                # Для скомпилированного приложения
                meipass_path = sys._MEIPASS
                content = os.listdir(meipass_path)

                message = f"Временная директория PyInstaller: {meipass_path}\n\n"
                message += "Содержимое временной директории:\n"
                message += "\n".join(content)
                messagebox.showinfo("Содержимое sys._MEIPASS", message)
            else:
                messagebox.showinfo("Информация", "Программа запущена из исходников, sys._MEIPASS не существует.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при отображении содержимого sys._MEIPASS: {str(e)}")

    # Вызов функции
    # show_meipass_content()

    def insert_images_to_word(self, doc):
        """Вставляет изображения из PDF в Word-документ на места меток"""
        try:
            if not self.pdf_var.get():
                logging.warning("PDF файл не выбран, пропускаем вставку изображений")
                return

            pdf_path = self.pdf_var.get()
            pdf_reader = PDFReader(pdf_path)
            pdf_reader.process_pdf()

            # Сопоставление типов графиков с метками
            image_mapping = {
                "Graph_1": "{{Picture3}}",  # Граф.Хорнера
                "Log_1": "{{Picture2}}",  # Диагностический (log)
                "Log_2": "{{Picture6}}",  # Диагностический сравнение
                "Semi_1": "{{Picture4}}",  # Полулогарифмический (semi),
                "History_1": "{{Picture1}}", # Обзорный на ВНК
                "History_2": "{{Picture7}}",  # Обзорный на ВНК сравнение
                "Map_1": "{{Picture5}}", # Карта
                "ACA_1": "{{Picture8}}",  # АСА график
            }

            # Замена меток изображениями
            for paragraph in doc.paragraphs:
                for image_type, placeholder in image_mapping.items():
                    if placeholder in paragraph.text:
                        image_path = os.path.join("plots", f"cropped_image_{image_type}.png")
                        if os.path.exists(image_path):
                            # Удаляем метку
                            paragraph.text = paragraph.text.replace(placeholder, "")

                            # Добавляем изображение
                            run = paragraph.add_run()
                            run.add_picture(image_path, width=Inches(6))  # Ширина 6 дюймов
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            logging.info(f"Вставлено изображение {image_type} на место {placeholder}")

            # Аналогичная замена в таблицах
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for image_type, placeholder in image_mapping.items():
                                if placeholder in paragraph.text:
                                    image_path = os.path.join("plots", f"cropped_image_{image_type}.png")
                                    if os.path.exists(image_path):
                                        paragraph.text = paragraph.text.replace(placeholder, "")
                                        run = paragraph.add_run()
                                        run.add_picture(image_path, width=Inches(4))  # Меньший размер для таблиц
                                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        except Exception as e:
            logging.error(f"Ошибка при вставке изображений: {str(e)}")
            raise RuntimeError(f"Не удалось вставить изображения: {str(e)}")

    def convert_docx_to_doc(self, docx_path, doc_path):
        """
        Конвертирует файл .docx в .doc с использованием Microsoft Word через pywin32.
        """
        word = None
        doc = None
        try:
            # Создаем объект Word
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # Открываем документ .docx
            doc = word.Documents.Open(docx_path)

            # Сохраняем как .doc
            doc.SaveAs(doc_path, FileFormat=0)  # FileFormat=0 означает формат Word 97-2003 (.doc)
            logging.info(f"Файл успешно сконвертирован: {doc_path}")

        except Exception as e:
            logging.error(f"Ошибка при конвертации .docx в .doc: {str(e)}")
            raise RuntimeError(f"Не удалось сконвертировать файл: {str(e)}")

        finally:
            # Закрываем документ и приложение Word
            if doc:
                doc.Close(SaveChanges=False)  # Явно указываем, что изменения не сохраняем
            if word:
                word.Quit()

            # Освобождаем ресурсы COM
            del doc
            del word

    def kill_excel_processes(self):
        """Принудительно закрывает все процессы Excel"""
        try:
            excel_app = win32com.client.Dispatch("Excel.Application")
            excel_app.Quit()
            del excel_app
        except Exception as e:
            logging.warning(f"Ошибка при закрытии Excel: {str(e)}")

        # Дополнительная проверка через taskkill
        try:
            os.system('taskkill /f /im excel.exe')
        except Exception as e:
            logging.warning(f"Ошибка при завершении процесса Excel: {str(e)}")

    def perform_calculations(self):
        """Выполняет расчеты для таблиц calculate, calculatedPressure и импорт предыдущих данных"""
        try:
            # Получаем ID основной записи
            main_data_id = self.ensure_main_data_exists()
            if main_data_id is None:
                messagebox.showerror("Ошибка", "Нет основной записи для расчета!")
                return

            # 0. Импорт данных предыдущего исследования
            try:
                # Получаем данные напрямую из базы, избегая проблем с pandas
                field_name = self.db.get_field_name(main_data_id)
                well_name = self.db.get_well_name(main_data_id)

                if field_name and well_name:
                    import_success = self.db.import_previous_research_data(main_data_id, field_name, well_name)
                    if import_success:
                        logging.info("Данные предыдущего исследования успешно импортированы")
                    else:
                        logging.info("Данные предыдущего исследования не найдены")
                else:
                    logging.warning(f"Недостаточно данных для импорта: field={field_name}, well={well_name}")

            except Exception as e:
                logging.warning(f"Не удалось импортировать предыдущие данные: {str(e)}")

            # 1. Расчет для calculatedPressure
            try:
                calculated_data = self.calculate_and_save_pressure_values(main_data_id)
                if calculated_data:
                    logging.info(f"Рассчитанные давления: {calculated_data}")
                else:
                    logging.warning("Не удалось рассчитать давления")
            except Exception as e:
                logging.error(f"Ошибка расчета давлений: {str(e)}")
                messagebox.showerror("Ошибка", f"Ошибка расчета давлений: {str(e)}")
                return

            # 2. Расчет для calculate таблицы
            try:
                self.db.update_calculate_table(main_data_id)
                logging.info("Таблица calculate успешно обновлена")
            except Exception as e:
                logging.error(f"Ошибка обновления таблицы calculate: {str(e)}")
                messagebox.showerror("Ошибка", f"Ошибка обновления таблицы calculate: {str(e)}")
                return

            messagebox.showinfo("Успех", "Расчеты успешно выполнены!\n\n"
                                         "• Импорт предыдущих данных (если найдены)\n"
                                         "• Рассчитаны давления (calculatedPressure)\n"
                                         "• Обновлена таблица calculate")

        except Exception as e:
            logging.error(f"Общая ошибка при выполнении расчетов: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при выполнении расчетов: {str(e)}")

    def calculate_and_save_pressure_values(self, main_data_id):
        """Рассчитывает и сохраняет значения для calculatedPressure"""
        try:

            # Вызываем расчет в базе данных
            calculated_data = self.db.calculate_pressure_values(main_data_id)

            if calculated_data:
                logging.info(f"Рассчитанные давления сохранены: {calculated_data}")
                return calculated_data
            else:
                logging.warning("Не удалось рассчитать давления")
                return None

        except Exception as e:
            logging.error(f"Ошибка расчета давлений: {str(e)}")
            raise

    def generate_report_logic(self, doc, output_file_path, selected_template):
        import win32com.client
        import os
        import docx
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        import pandas as pd
        from datetime import datetime
        import logging
        import pythoncom

        logging.info(f"Начало формирования отчета. Выходной файл: {output_file_path}, Шаблон: {selected_template}")
        try:
            logging.info(f"Output file path: {output_file_path}")
            logging.info(f"Selected template: {selected_template}")

            try:
                db = AccessDatabase()
                last_record = db.get_last_record()

                if last_record.empty:
                    raise ValueError("Нет данных для формирования отчета")

                main_data_id = last_record.iloc[0]['ID']

                # Получаем все данные из базы
                parameters = db.get_parameters(main_data_id)
                research_params = db.get_research_params(main_data_id)
                amendments = db.get_amendments(main_data_id)

                # Объединяем все данные в один словарь
                data = {}
                # Добавляем основные данные
                for column in last_record.columns:
                    value = last_record.iloc[0][column]
                    if pd.notna(value):
                        data[column.lower()] = value

                # Добавляем параметры, исследовательские параметры и поправки
                data.update(parameters)
                data.update(research_params)
                data.update(amendments)

                # Выбор шаблона Word
                template_map = {
                    "КВД_Заполярка": templates_path("KVD_Zapolyarka.docx"),
                    "КВД_Оренбург": templates_path("KVD_Orenburg.docx"),
                    "КВД_Оренбург_газ": templates_path("KVD_Orenburg_gas.docx"),
                    "КВД_Оренбург2": templates_path("KVD_Orenburg2.docx"),
                    "КВД_Хантос": templates_path("KVD_Khantos.docx"),
                    "КВД_глушение": templates_path("KVD_For_Killing.docx"),
                    "КВД_ННГ": templates_path("KVD_NNG.docx"),
                    "КВД+ИД": templates_path("KVD_ID.docx"),
                    "КСД": templates_path("KSD.docx"),
                    "КПД": templates_path("KPD.docx"),
                    "КПД+ИД": templates_path("KPD_ID.docx"),
                    "ГРП": templates_path("GRP.docx")
                }

                # Преобразуем ключи в нижний регистр для сравнения
                template_map_lower = {k.lower(): v for k, v in template_map.items()}

                if selected_template.lower() in template_map_lower:
                    selected_template_file = template_map_lower[selected_template.lower()]
                else:
                    raise ValueError(f"Неверное имя шаблона: {selected_template}. "
                                     f"Доступные шаблоны: {list(template_map.keys())}")

                # # Добавление данных из предыдущего исследования --------------------------------------
                try:
                    # Импортируем данные предыдущего исследования
                    field_name = data.get('field', '')
                    well_name = data.get('well', '')

                    if field_name and well_name:
                        success = self.db.import_previous_research_data(main_data_id, field_name, well_name)
                        if success:
                            logging.info("Данные предыдущего исследования успешно импортированы")
                        else:
                            logging.info("Данные предыдущего исследования не найдены или не требуются")

                except Exception as e:
                    logging.error(f"Ошибка при работе с историческими данными: {e}", exc_info=True)

                # После импорта данных получаем их для отчета
                previous_data = self.db.get_previous_research_data(main_data_id)
                if previous_data:
                    replace_tags_only(doc, previous_data)
                    logging.info("Данные предыдущего исследования добавлены в отчет")
                # try:
                #     field_name = data.get('field', '').replace(" ", "_").capitalize()
                #     previous_data_file = f'Итоговая таблица_{field_name}.xlsx'
                #     previous_data_path = table_prev_path(previous_data_file)
                #
                #     if os.path.exists(previous_data_path):
                #         try:
                #             final_table_df = pd.read_excel(previous_data_path, skiprows=11)
                #             well_num = data.get('well', '').split()[0] if data.get('well') else ''
                #             logging.info(f"Скважина: {well_num}")
                #
                #             # Фильтруем данные по номеру скважины
                #             final_table_df['Скважина'] = final_table_df['Скважина'].astype(str).str.strip()
                #             final_table_df = final_table_df.dropna(subset=['Скважина'])
                #             filtered_data = final_table_df[final_table_df['Скважина'] == well_num]
                #
                #             if not filtered_data.empty:
                #                 # Обработка данных из файла
                #                 pd.set_option('mode.use_inf_as_na', True)
                #                 filtered_data.loc[:, 'Дата испытания'] = filtered_data['Дата испытания'].apply(
                #                     lambda x: datetime.strptime(x, "%d.%m.%Y") if isinstance(x, str) else x
                #                 )
                #
                #                 latest_entry = filtered_data.loc[filtered_data['Дата испытания'].idxmax()]
                #
                #                 # СОХРАНЯЕМ ДАННЫЕ В ТАБЛИЦУ prevData
                #                 prev_data = {
                #                     'PplVnk': latest_entry.get('Рпл  на ВНК, кгс/см2'),
                #                     'PzabVnk': latest_entry.get('Рзаб  на ВНК, кгс/см2'),
                #                     'DataRes': latest_entry.get('Дата испытания'),
                #                     'Water': latest_entry.get('% воды'),
                #                     'Q': latest_entry.get('Qж/Qг, м3/сут'),
                #                     'Kprod': latest_entry.get('Кпрод. м3/сут*кгс/см2'),
                #                     'Smeh': latest_entry.get('Скин-фактор механич./интегр.'),
                #                     'Heff': latest_entry.get('Нэф., м.'),
                #                     'Kgidr': latest_entry.get('Кгидр., Д*см/сПз'),
                #                     'InputData_ID': main_data_id  # Связь с текущим исследованием
                #                 }
                #
                #                 # Сохраняем в базу данных
                #                 self.db.insert_prev_data(prev_data)
                #                 logging.info("Данные предыдущего исследования сохранены в базу")
                #
                #                 # Нормализуем строки для отчета
                #                 final_table_df = final_table_df.map(
                #                     lambda x: normalize_string(x) if isinstance(x, str) else x)
                #                 final_table_df.columns = [normalize_string(col) for col in final_table_df.columns]
                #
                #                 result_dict = {
                #                     normalize_string('Рпл  на ВНК, кгс/см2'): latest_entry['Рпл  на ВНК, кгс/см2'],
                #                     normalize_string('Рзаб  на ВНК, кгс/см2'): latest_entry['Рзаб  на ВНК, кгс/см2'],
                #                     normalize_string('Дата испытания'): latest_entry['Дата испытания'].strftime(
                #                         "%d.%m.%Y"),
                #                     normalize_string('% воды'): str(latest_entry['% воды']),
                #                     normalize_string('Qж/Qг, м3/сут   '): str(latest_entry['Qж/Qг, м3/сут   ']),
                #                     normalize_string('Кпрод. м3/сут*кгс/см2'): str(
                #                         latest_entry['Кпрод. м3/сут*кгс/см2']),
                #                     normalize_string('Скин-фактор механич./интегр.'): str(
                #                         latest_entry['Скин-фактор механич./интегр.']),
                #                     normalize_string('Нэф., м.'): str(latest_entry['Нэф., м. ']),
                #                     normalize_string('Кгидр., Д*см/сПз'): str(latest_entry['Кгидр., Д*см/сПз'])
                #                 }
                #
                #                 replace_tags_only(doc, result_dict)
                #                 logging.info("Данные из файла предыдущих исследований успешно загружены.")
                #
                #         except Exception as e:
                #             logging.error(f"Ошибка при чтении файла предыдущих данных Excel: {str(e)}")
                #
                #     else:
                #         logging.warning(f"Файл предыдущих данных не найден: {previous_data_path}")
                #         logging.info("Для данного отчета не требуется файл с предыдущими данными.")
                #
                # except Exception as e:
                #     logging.error(f"Ошибка при работе с историческими данными: {e}", exc_info=True)

                # Расчет плотности
                KVD_density = data.get('dens2', 0)
                work_density = data.get('dens1', 0)
                if not work_density:
                    density = f'{KVD_density} г/см3'
                else:
                    density = f'{KVD_density} г/см3 для пересчета участка КВД и {work_density} г/см3 - для пересчета цикла отработки скважины'

                # Добавляем density в данные
                data['density'] = density

                # Расчет Leff1 в зависимости от модели
                model_value = data.get('model', '')
                if model_value == "Горизонтальная с ГРП":
                    Leff1 = round(data.get('Leff1', 0))
                else:
                    Leff1 = round(data.get('Leff1', 0))

                # Добавляем Leff1 в данные
                data['Leff1'] = Leff1

                model_params = {
                    "Вертикальная": {
                        "additional_params": [],
                        "diagnostic_text": "model_descriptions.Вертикальная"
                    },
                    "Наклонн.": {
                        "additional_params": [],
                        "diagnostic_text": "model_descriptions.Наклонн."
                    },
                    "Вертикальная - частичное вскрытие": {
                        "additional_params": [
                            {"name": "Скин-фактор механический", "value": data.get("S_meh1"), "key": "S_meh1"},
                            {"name": "Скин-фактор геометрический", "value": data.get("S_geom1"), "key": "S_geom1"},
                            {"name": "Эффективная часть интервала перфорации (hw), (м)", "value": data.get("Leff1"),
                             "key": "Leff1"}
                        ],
                        "diagnostic_text": "model_descriptions.Вертикальная - частичное вскрытие"
                    },
                    "Горизонтальн.": {
                        "additional_params": [
                            {"name": "Скин-фактор механический", "value": data.get("S_meh1"), "key": "S_meh1"},
                            {"name": "Эффективная длина скважины, (м)", "value": data.get("Leff1"), "key": "Leff1"}
                        ],
                        "diagnostic_text": "model_descriptions.Горизонтальн."
                    },
                    "Горизонтальная с ГРП": {
                        "additional_params": [
                            {"name": "Скин-фактор механический", "value": data.get("S_meh1"), "key": "S_meh1"},
                            {"name": "Количество трещин", "value": data.get("num_frac1"), "key": "num_frac1"},
                            {"name": "Полудлина трещины, (м)", "value": data.get("Xf1"), "key": "Xf1"}
                        ],
                        "diagnostic_text": "model_descriptions.Горизонтальная с ГРП"
                    },
                    "Трещина - бесконечная проводимость": {
                        "additional_params": [
                            {"name": "Скин кольматации стенок трещины", "value": data.get("S_meh1"), "key": "S_meh1"},
                            {"name": "Полудлина трещины, (м)", "value": data.get("Xf1"), "key": "Xf1"}
                        ],
                        "diagnostic_text": "model_descriptions.Трещина - бесконечная проводимость"
                    },
                    "Трещина - конечная проводимость": {
                        "additional_params": [
                            {"name": "Скин кольматации стенок трещины", "value": data.get("S_meh1"), "key": "S_meh1"},
                            {"name": "Полудлина трещины, (м)", "value": data.get("Xf1"), "key": "Xf1"},
                            {"name": "Проводимость трещины, (Fc)", "value": data.get("Fc1"), "key": "Fc1"}
                        ],
                        "diagnostic_text": "model_descriptions.Трещина - конечная проводимость"
                    },
                    "Трещина - равномерный поток": {
                        "additional_params": [
                            {"name": "Скин кольматации стенок трещины", "value": data.get("S_meh1"), "key": "S_meh1"},
                            {"name": "Полудлина трещины, (м)", "value": data.get("Xf1"), "key": "Xf1"}
                        ],
                        "diagnostic_text": "model_descriptions.Трещина - равномерный поток"
                    }
                }

                # Вставка параметров модели в таблицу результатов
                def insert_model_params_to_table(doc, model_name, data):
                    """Вставляет параметры модели после строки 'Проницаемость, (мД)'"""
                    table = find_results_table(doc)
                    if not table:
                        logging.error("Таблица результатов не найдена")
                        return False

                    params = model_params.get(model_name, {}).get("additional_params", [])
                    if not params:
                        logging.warning(f"Нет дополнительных параметров для модели {model_name}")
                        return False

                    # Находим индекс строки с "Проницаемость, (мД)"
                    target_row_idx = -1
                    for i, row in enumerate(table.rows):
                        if row.cells and "Проницаемость, (мД)" in row.cells[0].text:
                            target_row_idx = i
                            break

                    if target_row_idx == -1:
                        logging.error("Строка 'Проницаемость, (мД)' не найдена в таблице")
                        return False

                    # Вставляем параметры после найденной строки
                    for param in params:
                        new_row = table.add_row()
                        table.rows._tbl.insert(target_row_idx + 1, new_row._tr)
                        target_row_idx += 1

                        # Форматируем и вставляем данные
                        name = format_units(param["name"])
                        value = format_units(str(param["value"]))

                        new_row.cells[0].text = name
                        new_row.cells[1].text = value

                        # Устанавливаем выравнивание
                        for paragraph in new_row.cells[0].paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            set_font_size(paragraph, 12)

                        for paragraph in new_row.cells[1].paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            set_font_size(paragraph, 12)

                    logging.info(f"Успешно вставлено {len(params)} параметров модели")
                    return True

                # Основной блок обработки документа
                model_value = data.get('model', '')
                model_name = None

                # Находим соответствующую модель
                for key in model_params.keys():
                    if model_value and key.strip().lower() in model_value.strip().lower():
                        model_name = key
                        break

                if model_name is None:
                    model_name = "Вертикальная"
                    logging.warning(
                        f"Предупреждение: Модель '{model_value}' не найдена в model_params. Используется модель по умолчанию: {model_name}")

                # Добавляем параметры модели в словарь data
                model_data = model_params.get(model_name, {})
                model_description = TEXT_TEMPLATES["model_descriptions"].get(model_name, "")

                if not model_description:
                    logging.warning(f"Описание для модели '{model_name}' не найдено в шаблонах")

                data.update({
                    "model_description": model_description,
                    "model_name": model_name,
                    "diagnostic_text": model_description,
                    **{f"param_{k}": v for k, v in enumerate(model_data.get("additional_params", []))},
                })

                replace_tags_only(doc, data)
                logging.info("Метки в отчете успешно заменены на значения.")

                # Специальная обработка diagnostic_text
                for paragraph in doc.paragraphs:
                    if "{{diagnostic_text}}" in paragraph.text:
                        if paragraph.runs:
                            original_font = paragraph.runs[0].font

                        paragraph.text = paragraph.text.replace("{{diagnostic_text}}", data["model_description"])

                        if paragraph.runs and original_font:
                            paragraph.runs[0].font.name = original_font.name
                            paragraph.runs[0].font.size = original_font.size
                            paragraph.runs[0].font.bold = original_font.bold
                            paragraph.runs[0].font.italic = original_font.italic
                            paragraph.runs[0].font.underline = original_font.underline
                            if original_font.color.rgb:
                                paragraph.runs[0].font.color.rgb = original_font.color.rgb

                # Вставка параметров модели в таблицу
                if not insert_model_params_to_table(doc, model_name, data):
                    logging.warning("Не удалось вставить параметры модели в таблицу")

                # Проверка и сохранение
                if "{{diagnostic_text}}" in [p.text for p in doc.paragraphs]:
                    logging.error("Метка diagnostic_text не была заменена!")
                else:
                    logging.info("метка diagnostic_text успешно заменена")

                # Удаление лишних строк из таблицы результатов
                replace_and_format_table(doc, data)

                doc.save(output_file_path)
                logging.info(f"Data dictionary content: {json.dumps(data, indent=2, ensure_ascii=False)}")

                # Внесение данных в Helper----------------------------------------------------------
                try:
                    logger.info("Начало обновления Helper.xlsm")

                    if data.get("type_of_research") in ["КВД", "КВУ", "КСД", "ИК+КВУ", "ИК+КВД", "ГРП", "ГДП",
                                                        "ИД+КВД"]:
                        category = "доб"
                    else:
                        category = "нагн"

                    new_row = {
                        "Дата интерпретации": data.get("date_of_analiz"),
                        "Дата начала исследования": data.get("date_research"),
                        "Дата конца исследования": data.get("date_researcf"),
                        "ДО": data.get("company"),
                        "Месторождение": data.get("field"),
                        "Пласт": data.get("formation"),
                        "Куст": data.get("well", "").split()[2] if len(data.get("well", "").split()) > 2 else data.get(
                            "well", ""),
                        "№скв.": data.get("well", "").split()[0] if data.get("well") else data.get("well", ""),
                        "Категория скважин": category,
                        "Вид исследования": data.get("type_of_research"),
                        "Исполнитель (организация)": "ИТС",
                        "Интерпретатор": data.get("interpreter"),
                        "Наличие в базе": "база",
                        "Оборудование": data.get("device", "").split()[1] if len(
                            data.get("device", "").split()) > 1 else data.get("device", ""),
                        "Назначение": "Запрос ДО",
                        "Класс исследования": data.get("klass"),
                        "Успешность": data.get("success"),
                        "Длительность факт": data.get("duration"),
                        "Qн": data.get("Qoil")
                    }

                    # Открытие файла Excel
                    logging.info("Обновление Helper.xlsm")
                    excel_helper = win32com.client.DispatchEx("Excel.Application")
                    excel_helper.Visible = False
                    excel_helper.DisplayAlerts = False

                    excel_file_path_helper = resource_path("Helper.xlsm")
                    if not os.path.exists(excel_file_path_helper):
                        logger.warning(f"Файл Helper.xlsm не найден: {excel_file_path_helper}")
                        return False

                    # Открываем Helper.xlsm
                    workbook = excel_helper.Workbooks.Open(excel_file_path_helper)
                    sheet1 = workbook.Sheets['Sheet1']

                    # Находим первую полностью пустую строку
                    empty_row_index = sheet1.Cells(sheet1.Rows.Count, 4).End(-4162).Row + 1

                    # Добавляем значения из словаря в соответствующие столбцы
                    for key, value in new_row.items():
                        column_index = None
                        for col in sheet1.Range(sheet1.Cells(1, 1), sheet1.Cells(1, sheet1.Columns.Count)).Columns:
                            if col.Cells(1, 1).Value == key:
                                column_index = col.Cells(1, 1).Column
                                break

                        if column_index is not None:
                            sheet1.Cells(empty_row_index, column_index).Value = value

                    # Сохраняем книгу
                    workbook.Save()
                    logger.info("Данные успешно записаны в Helper.xlsm")

                except Exception as e:
                    logger.error(f"Ошибка при работе с Helper.xlsm: {str(e)}")
                finally:
                    if 'workbook' in locals() and workbook is not None:
                        try:
                            workbook.Close(SaveChanges=True)
                        except Exception as e:
                            logging.warning(f"Ошибка при закрытии Workbook: {str(e)}")
                    if excel_helper:
                        excel_helper.Quit()

                logger.info("Отчет успешно сформирован!")
                return True

            except Exception as e:
                logging.error(f"Ошибка при формировании отчета: {str(e)}", exc_info=True)
                raise RuntimeError(f"Ошибка при формировании отчета: {str(e)}")

            finally:
                if db:
                    db.close()

        except Exception as e:
            import traceback
            traceback.print_exc()
            logging.error(f"Ошибка при формировании отчета: {str(e)}", exc_info=True)
            return False


    def generate_report(self):
        """УПРОЩЕННАЯ ВЕРСИЯ ДЛЯ РАБОТЫ С БАЗОЙ ДАННЫХ"""
        global db
        try:
            logging.info("Начало формирования отчета")

            if not hasattr(self, 'output_directory') or not self.output_directory:
                logging.error("Не выбрана папка для сохранения отчета")
                messagebox.showerror("Ошибка", "Сначала выберите папку для сохранения!")
                return False

            # Получаем данные из базы
            db = AccessDatabase()
            last_record = db.get_last_record()

            if last_record.empty:
                messagebox.showerror("Ошибка", "Нет данных для формирования отчета")
                return False

            main_data_id = last_record.iloc[0]['id']
            parameters = db.get_parameters(main_data_id)
            research_params = db.get_research_params(main_data_id)
            amendments = db.get_amendments(main_data_id)

            # Объединяем все данные
            import pandas as pd  # Добавляем импорт
            data = {}
            for column in last_record.columns:
                value = last_record.iloc[0][column]
                if pd.notna(value):
                    data[column.lower()] = value
            data.update(parameters)
            data.update(research_params)
            data.update(amendments)

            template_mapping = {
                "КВД_Заполярка": templates_path("KVD_Zapolyarka.docx"),
                "КВД_Оренбург": templates_path("KVD_Orenburg.docx"),
                "КВД_Оренбург_газ": templates_path("KVD_Orenburg_gas.docx"),
                "КВД_Оренбург2": templates_path("KVD_Orenburg2.docx"),
                "КВД_Хантос": templates_path("KVD_Khantos.docx"),
                "КВД_глушение": templates_path("KVD_For_Killing.docx"),
                "КВД_ННГ": templates_path("KVD_NNG.docx"),
                "КВД+ИД": templates_path("KVD_ID.docx"),
                "КСД": templates_path("KSD.docx"),
                "КПД": templates_path("KPD.docx"),
                "КПД+ИД": templates_path("KPD_ID.docx"),
                "ГРП": templates_path("GRP.docx")
            }

            # Получаем выбранный шаблон (ключ)
            selected_template_key = self.template_var.get().strip()
            print(f"Выбранный шаблон (ключ): {selected_template_key}")

            # Проверяем, существует ли ключ в словаре
            if selected_template_key not in template_mapping:
                messagebox.showerror("Ошибка", f"Шаблон '{selected_template_key}' не найден!")
                return False

            if selected_template_key.lower() == "квд_глушение":
                # Создаем временный файл KVD_For_Killing.docx
                logging.info("Создание файла KVD_For_Killing.docx...")
                result = copy_excel_to_word_pandas(
                    excel_path=resource_path('Report.xlsx'),
                    word_path=os.path.abspath(templates_path('КВД для глушения_prev.docx')),
                    sheet_name='current',
                    search_text='Prognoz_Ppl'
                )

                if result is None:
                    logging.error("Ошибка: copy_excel_to_word_pandas вернул None")
                    raise RuntimeError("Не удалось создать временный Word-документ")

                logging.info(f"Файл KVD_For_Killing.docx успешно создан: {result}")

                # Обновляем имя шаблона
                selected_template_file = os.path.basename(result)
            else:
                # Получаем имя файла шаблона из словаря
                selected_template_file = template_mapping[selected_template_key]
                logging.info(f"Имя файла шаблона: {selected_template_file}")

            # Формируем полный путь к файлу шаблона
            template_path = resource_path(selected_template_file)
            print(f"Путь к шаблону: {template_path}")
            logging.info(f"Путь к шаблону Word: {template_path}")

            # Проверяем существование файла шаблона
            if not os.path.exists(template_path):
                logging.error(f"Шаблон не найден: {template_path}")
                messagebox.showerror("Ошибка", f"Шаблон '{selected_template_file}' не найден!")
                return False

            # Создаем объект Document из шаблона
            doc = Document(template_path)

            # Функция для получения значений из данных БД вместо Excel
            def get_value_from_data(key, default="Без_данных"):
                """Получает значение из объединенных данных БД"""
                value = data.get(key, default)
                if isinstance(value, (datetime, pd.Timestamp)):
                    return value.strftime("%d.%m.%Y")
                elif isinstance(value, str):
                    # Пытаемся преобразовать строку в дату, если это возможно
                    try:
                        date_obj = datetime.strptime(value, "%d.%m.%Y")
                        return date_obj.strftime("%d.%m.%Y")
                    except (ValueError, AttributeError):
                        return value.strip() if value else default
                return str(value) if value else default

            # Формируем базовое имя файла
            base_name = (
                f"Закл_"
                f"{get_value_from_data('type_of_research', 'Без_данных')}_"
                f"{get_value_from_data('field', 'Без_данных')}_"
                f"{get_value_from_data('well', 'Без_данных')}_"
                f"{get_value_from_data('date_research', 'Без_данных')}"
            )

            # Убираем недопустимые символы
            import re
            clean_name = re.sub(r'[<>:"/\\|?*]', '_', base_name)
            report_name_docx = f"{clean_name}.docx"
            report_name_doc = f"{clean_name}.doc"

            # Полный путь для сохранения
            output_file_path_docx = os.path.normpath(os.path.join(self.output_directory, report_name_docx))
            output_file_path_doc = os.path.normpath(os.path.join(self.output_directory, report_name_doc))

            # Проверяем и обрабатываем дубликаты
            counter = 1
            while os.path.exists(output_file_path_docx) or os.path.exists(output_file_path_doc):
                new_name = f"{clean_name}_{counter}"
                output_file_path_docx = os.path.join(self.output_directory, f"{new_name}.docx")
                output_file_path_doc = os.path.join(self.output_directory, f"{new_name}.doc")
                counter += 1

            # Вставка изображений
            self.insert_images_to_word(doc)

            # Сохранение временного .docx файла
            doc.save(output_file_path_docx)
            logging.info(f"Отчет успешно сохранен во временном формате: {output_file_path_docx}")

            # Вызываем функцию для формирования отчета
            logging.info(f"Вызов generate_report_logic с параметрами: {output_file_path_docx}, {selected_template_key}")
            success = self.generate_report_logic(doc, output_file_path_docx, selected_template_key)
            fix_units(doc)
            if success:
                # Сохраняем финальную версию документа
                doc.save(output_file_path_docx)
                # Конвертация .docx в .doc
                self.convert_docx_to_doc(output_file_path_docx, output_file_path_doc)

                logging.info(f"Отчет успешно сохранен: {output_file_path_doc}")
                messagebox.showinfo("Успех", f"Отчет сформирован успешно: {output_file_path_doc}")

                # Удаление временного файла
                time.sleep(3)  # Ждем, пока файл перестанет быть заблокированным
                try:
                    if os.path.exists(output_file_path_docx):
                        os.remove(output_file_path_docx)
                        logging.info(f"Временный файл удален: {output_file_path_docx}")
                except Exception as e:
                    logging.warning(f"Не удалось удалить временный файл: {str(e)}")
            else:
                logging.error("Не удалось сформировать отчет")
                messagebox.showerror("Ошибка", "Не удалось сформировать отчет.")
                return False

        except Exception as e:
            logging.error(f"Ошибка при формировании отчета: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при формировании отчета: {str(e)}")
            return False
        finally:
            if 'db' in locals():
                db.close()


if __name__ == "__main__":
    root = tk.Tk()
    app = ReportGUI(root)
    root.mainloop()
